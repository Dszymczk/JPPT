import os
from docx import Document

Filename = "Rozdzial"

tmp_list = os.listdir()
files_list = []
for file in tmp_list:
    if ".txt" in file:
        files_list.append(file)

document = Document()
for file in files_list:
    document.add_heading(file, 0)
    f = open(file, "r")
    text = f.read()
    document.add_paragraph(text)

document.save('Tekst.docx')
